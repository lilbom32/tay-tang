#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Build HDV Workflow Docx for Tour Tay Tang 9N8D
Integrates: CFM info, THUYET_MINH_CHI_TIET_9_NGAY, MOMENTS_OF_TRUTH,
ROUTE_DIAGRAMS, VISUAL_GUIDE, and encyclopedia content from Tây Tạng.docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def set_run_font(run, font_name='Times New Roman', size=11, bold=False, italic=False, color=None):
    font = run.font
    font.name = font_name
    font.size = Pt(size)
    font.bold = bold
    font.italic = italic
    if color:
        font.color.rgb = RGBColor(*color)
    # Set East Asian font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.insert(0, rFonts)

def add_heading_custom(doc, text, level=1, font_name='Times New Roman', color=(0, 51, 102)):
    """Add a styled heading."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    sizes = {1: 16, 2: 14, 3: 12}
    set_run_font(run, font_name, sizes.get(level, 11), bold=True, color=color)
    p.space_after = Pt(6 if level > 1 else 10)
    p.space_before = Pt(12 if level > 1 else 16)
    return p

def add_para_custom(doc, text, bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT,
                    font_name='Times New Roman', first_line_indent=None, space_after=6):
    p = doc.add_paragraph()
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, font_name, size, bold=bold, italic=italic)
    p.alignment = align
    return p

def add_bullet(doc, text, level=0, size=11, font_name='Times New Roman'):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    run = p.runs[0] if p.runs else p.add_run(text)
    if not p.runs:
        run.text = text
    set_run_font(run, font_name, size)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p

def add_table_custom(doc, headers, rows, col_widths=None, header_color='D9E2F3'):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], header_color)
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, 'Times New Roman', 10, bold=True)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            if i < len(row_cells):
                row_cells[i].text = str(val)
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, 'Times New Roman', 10)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(w)
    doc.add_paragraph()
    return table

def add_image_with_caption(doc, img_path, caption, width=Inches(6)):
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=width)
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_p.add_run(f"[Hình] {caption}")
        set_run_font(cap_run, 'Times New Roman', 10, italic=True, color=(100,100,100))
        doc.add_paragraph()
    else:
        add_para_custom(doc, f"[Hình tham khảo: {caption} — đặt tại {img_path}]", italic=True, size=10, color=(150,150,150))

# ============================================
# MAIN DOCUMENT BUILDER
# ============================================
doc = Document()

# Set default font for document
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
style.paragraph_format.space_after = Pt(6)

# ============================================
# TITLE PAGE
# ============================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CẨM NANG VẬN HÀNH & THUYẾT MINH HDV")
set_run_font(run, 'Times New Roman', 22, bold=True, color=(0, 51, 102))
p.space_after = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Tour Tây Tạng Huyền Bí 9N8Đ | Lhasa · Shigatse · Everest Base Camp")
set_run_font(run, 'Times New Roman', 14, bold=True, color=(80,80,80))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Khởi hành: 26/04/2026 | Đối tác: Yunnan DiLang International Travel Agency Co., Ltd")
set_run_font(run, 'Times New Roman', 11, italic=True, color=(100,100,100))
doc.add_paragraph()

# Insert CFM render page 1 as cover visual
if os.path.exists('_render_cfm/page-1.png'):
    add_image_with_caption(doc, '_render_cfm/page-1.png', "CFM – Thông tin cơ bản tour (trang 1)", width=Inches(6.2))

# ============================================
# PHẦN 1: TỔNG QUAN WORKFLOW
# ============================================
add_heading_custom(doc, "PHẦN I: TỔNG QUAN WORKFLOW & NGUYÊN TẮC HDV", level=1)

add_heading_custom(doc, "1. Flow cảm xúc 9 ngày", level=2)
flow_rows = [
    ["1", "Hạ cánh trên nóc nhà", "3.650m", "Xa lạ, hồi hộp, tôn nghiêm", "Địa lý – Sinh học – Tâm linh nhập môn"],
    ["2", "Trước khi có Phật giáo", "3.500–3.700m", "Tò mò, khai sáng", "Tiền sử – Bön – Nhập Phật"],
    ["3", "Đường lên trời", "4.400–5.020m", "Choáng ngợp, khiêm nhường", "Thủy văn – Khí hậu – Biến đổi khí hậu"],
    ["4", "Đống Vinh Quang", "3.800–5.000m", "Trầm lắng, nghiêm trang", "Gelug – Panchen – Tượng học"],
    ["5", "Dưới chân Nữ Thần", "5.200m", "Đỉnh điểm hùng vĩ", "Himalaya – Thần thoại – Leo núi"],
    ["6", "Dòng sông Mẹ", "3.650–4.000m", "Lắng đọng, hoài niệm", "Dân tộc học – Nông nghiệp – Huyền sử"],
    ["7", "Trái tim Lhasa", "3.650m", "Sâu lắng, đồng cảm", "Hành hương – Tín ngưỡng hàng ngày"],
    ["8", "Vào trong đá", "3.650–4.880m", "Tĩnh lặng, sáng tạo", "Thiền định – Thangka – Nghệ thuật"],
    ["9", "Mang theo Tây Tạng", "0m (về nhà)", "Biết ơn, mở mang", "Tổng hợp – Phản tỉnh – Hành động"],
]
add_table_custom(doc, ["Ngày", "Chủ đề chính", "Độ cao TB", "Cảm xúc chủ đạo", "Layer tri thức chính"], flow_rows, col_widths=[0.6, 1.8, 1.0, 1.6, 2.4])

add_heading_custom(doc, "2. Nguyên tắc vàng cho HDV (Moments of Truth)", level=2)
principles = [
    "Im lặng là công cụ thuyết minh mạnh nhất — khoảnh khắc không lời tạo ra tác động sâu hơn bất kỳ câu nào.",
    "Để cảnh quan nói trước, HDV nói sau — không bao giờ thuyết minh trước khi khách có ít nhất 60–90 giây phản ứng tự nhiên.",
    "Câu hỏi tốt hơn câu trả lời — một câu hỏi đúng lúc mở ra suy nghĩ sâu hơn 10 phút thuyết minh.",
    "Không xin lỗi khi thời tiết xấu — không có thời tiết xấu, chỉ có góc nhìn chưa đúng.",
    "Cảm xúc lan truyền — trạng thái nội tâm HDV là một phần của sản phẩm tour.",
    "HDV vào sau cùng, ra trước — tại các điểm tác động thị giác mạnh, để khách vào trước.",
    "Không đứng giữa khách và cảnh quan — tư thế vật lý quan trọng.",
    "Không bao giờ kiểm tra điện thoại trong MOT — điện thoại phải tắt hoặc im lặng hoàn toàn.",
    "Reframe thay vì xin lỗi — khi điều gì đó không theo kế hoạch, tìm góc nhìn khác.",
    "Biết khi nào không cần nói gì thêm — khi một khoảnh khắc đã thành công, đừng thêm bất cứ điều gì.",
]
for pr in principles:
    add_bullet(doc, pr)
doc.add_paragraph()

add_heading_custom(doc, "3. Bảng tóm tắt 7 Moments of Truth", level=2)
mot_rows = [
    ["1", "Lễ Khata", "Ngày 1", "Sân bay Gonggar", "HDV Tạng trao từng người, không vội", "Làm bên trong sảnh, như phát tờ rơi"],
    ["2", "Im lặng trước hồ Yamdrok", "Ngày 3", "Đèo Kampa La", "90 giây im lặng TRƯỚC khi thuyết minh", "Nói trước khi khách có phản ứng cảm xúc"],
    ["3", "Tượng Jampa Di Lặc", "Ngày 4", "Tashilhunpo", "60 giây khách tự phản ứng, HDV vào sau", "Thuyết minh ngay khi khách vừa vào cửa"],
    ["4", "Nhìn thấy Chomolungma", "Ngày 4", "Vùng Tingri", "Reframe tên: 'Không phải Everest'", "Nói 'tiếc là không thấy' khi trời mây"],
    ["5", "Đứng ở EBC 5.200m", "Ngày 5", "EBC North", "60 giây im lặng, HDV ngồi xuống cùng", "Lấp đầy im lặng bằng số liệu leo núi"],
    ["6", "Gặp Jowo Shakyamuni", "Ngày 7", "Đền Jokhang", "Không thuyết minh bên trong điện", "Giải thích iconography trong chánh điện"],
    ["7A", "Hoàng hôn Potala", "Ngày 8", "Quảng trường Potala", "Câu chuyện 'ngôi nhà trống' 65 năm", "Thúc giục về ăn tối, xe nổ máy chờ"],
]
add_table_custom(doc, ["#", "Tên MOT", "Ngày", "Địa điểm", "Yếu tố quyết định", "Failure Mode"], mot_rows, col_widths=[0.4, 1.3, 0.6, 1.3, 1.8, 1.6])

add_heading_custom(doc, "4. Bản đồ độ cao toàn chuyến", level=2)
if os.path.exists('_render_cfm/page-2.png'):
    add_image_with_caption(doc, '_render_cfm/page-2.png', "Profile độ cao 9 ngày – từ 3.650m lên 5.200m", width=Inches(6.0))

add_heading_custom(doc, "5. Bảng Oxy & Triệu chứng cần theo dõi", level=2)
oxy_rows = [
    ["0–1.000m", "100–90%", "Hoàn toàn bình thường", "—"],
    ["1.000–2.000m", "90–82%", "Bình thường, hơi mệt khi gắng sức", "—"],
    ["2.000–3.000m", "82–75%", "Hơi mệt khi leo dốc nhanh", "Bắt đầu uống nhiều nước"],
    ["3.000–3.650m (Lhasa)", "75–65%", "Đau đầu nhẹ, mệt mỏi, khó ngủ, tiểu nhiều", "Nghỉ ngơi ngay, uống nước"],
    ["3.650–4.500m", "65–58%", "Đau đầu nhiều hơn, hơi thở gấp khi đi nhanh", "Đi chậm, dừng thường xuyên"],
    ["4.500–5.000m", "58–53%", "Khó thở khi đi bộ bình thường, mệt rõ rệt", "Không gắng sức, quan sát liên tục"],
    ["5.000–5.200m (EBC)", "53–50%", "Mỗi bước đi cảm nhận được rõ ràng", "Giới hạn thời gian tối đa 90 phút"],
]
add_table_custom(doc, ["Độ cao", "Oxy so với mực biển", "Cảm giác điển hình", "Dấu hiệu cần chú ý"], oxy_rows, col_widths=[1.2, 1.2, 1.8, 1.8])

add_heading_custom(doc, "6. Nhận biết Hội chứng Nguy hiểm – Phải xuống ngay", level=2)
ams_rows = [
    ["AMS", "Acute Mountain Sickness", "Đau đầu + buồn nôn + mệt mỏi chung", "Nghỉ ngơi, uống nước, Diamox 125mg nếu có; không leo tiếp"],
    ["HACE ⚠️", "High Altitude Cerebral Edema", "Mất thăng bằng, lú lẫn, đau đầu dữ dội, nôn liên tục", "XUỐNG NGAY, gọi cấp cứu, oxy bổ sung nếu có"],
    ["HAPE ☠️", "High Altitude Pulmonary Edema", "Khó thở ngay cả khi nằm nghỉ, môi/móng tay tím tái, khò khè", "XUỐNG NGAY – NGUY HIỂM TÍNH MẠNG, gọi 120 ngay lập tức"],
]
add_table_custom(doc, ["Hội chứng", "Tên đầy đủ", "Triệu chứng nhận biết", "Xử lý"], ams_rows, col_widths=[0.8, 1.4, 2.2, 2.0], header_color='F2D7D5')

doc.add_page_break()

# ============================================
# PHẦN 2: KIẾN THỨC NỀN TÂY TẠNG
# ============================================
add_heading_custom(doc, "PHẦN II: KIẾN THỨC NỀN TÂY TẠNG — DỮ LIỆU THAM CHIẾU NHANH", level=1)

add_heading_custom(doc, "1. Vị trí & Quy mô", level=2)
geo_rows = [
    ["Độ cao trung bình", ">4.000–4.500 mét (13.000–14.800 ft)"],
    ["Diện tích", "~2,5 triệu km² — gấp gần 4 lần bang Texas"],
    ["Vị trí", "Giữa Trung Quốc phía bắc/đông, Ấn Độ, Nepal, Bhutan phía nam"],
    ["Tọa độ", "26°00'–39°47' Vĩ độ Bắc, 73°19'–104°47' Kinh độ Đông"],
    ["Biên giới tự nhiên", "Himalaya phía nam, núi Kunlun phía bắc, dãy Tanggula, Karakoram"],
]
add_table_custom(doc, ["Đặc điểm", "Chi tiết"], geo_rows, col_widths=[1.5, 3.5])

add_heading_custom(doc, "2. Ba vùng địa lý chính", level=2)
zone_rows = [
    ["Bắc Tây Tạng (Changthang)", "Thảo nguyên hoang mạc lạnh, đất đóng băng vĩnh cửu, hồ muối, ít dân"],
    ["Nam Tây Tạng (Thung lũng sông Yarlung Tsangpo)", "Vùng nông nghiệp chính, khí hậu ôn hòa hơn, trung tâm văn hóa"],
    ["Đông Tây Tạng (Kham)", "Núi cao hiểm trở, hẻm sông sâu, đa dạng sinh học cao"],
]
add_table_custom(doc, ["Vùng", "Đặc điểm"], zone_rows, col_widths=[1.8, 3.2])

add_heading_custom(doc, "3. Đỉnh cao huyền thoại", level=2)
peak_rows = [
    ["Everest (Chomolangma)", "8.848,86m", "Đỉnh cao nhất Trái Đất"],
    ["K2", "8.611m", "Đỉnh cao thứ hai thế giới"],
    ["Kailash", "6.638m", "Núi thiêng 4 tôn giáo, cấm leo"],
    ["Namcha Barwa", ">7.700m", "Đỉnh phía đông Himalaya, chuyển tiếp từ cao nguyên khô hạn sang rừng nhiệt đới"],
]
add_table_custom(doc, ["Đỉnh núi", "Độ cao", "Đặc biệt"], peak_rows, col_widths=[1.5, 1.0, 2.5])

add_heading_custom(doc, "4. Các hồ lớn", level=2)
lake_rows = [
    ["Namtso", "4.718m", "Hồ muối cao nhất thế giới, 'Hồ Thiên Đường'"],
    ["Yamdrok", "4.441m", "Hồ ngọc bích thiêng liêng — điểm dừng Ngày 3"],
    ["Manasarovar", "4.590m", "Hồ linh thiêng bậc nhất, gần núi Kailash"],
]
add_table_custom(doc, ["Hồ", "Độ cao", "Đặc điểm"], lake_rows, col_widths=[1.2, 1.0, 2.8])

add_heading_custom(doc, "5. Khí hậu theo vùng", level=2)
climate_rows = [
    ["Đông Nam (Nyingchi)", "Ẩm ướt, ôn đới, rừng rậm — 'Thụy Sĩ của Tây Tạng'"],
    ["Trung tâm (Lhasa, Shigatse)", "Khô, nắng, chênh lệch nhiệt lớn"],
    ["Tây (Ngari)", "Hoang mạc lạnh, gió mạnh, khô cằn"],
    ["Bắc (Nagqu)", "Thảo nguyên băng giá, bão cát, lạnh sâu"],
]
add_table_custom(doc, ["Vùng", "Khí hậu"], climate_rows, col_widths=[1.5, 3.5])

add_heading_custom(doc, "6. Tôn giáo — Bốn tông phái chính", level=2)
sect_rows = [
    ["Nyingma (Cổ phái / Hồng giáo)", "Truyền thừa từ thởi đế chế Tubo, Padmasambhava là tổ sư. Dzogchen (Đại viên mãn) là cốt lõi."],
    ["Sakya (Tạp giáo / Hoa giáo)", "Dòng họ Khön, nổi tiếng về học thuật uyên bác và nghi thuật Mật tông. Từng cai trị Tây Tạng thế kỷ 13."],
    ["Kagyu (Khách giáo / Bạch giáo)", "Khẩu truyền từ thầy đến trò. Mahamudra (Đại ấn) là cốt lõi. Milarepa là biểu tượng."],
    ["Gelug (Hoàng giáo)", "Tsongkhapa sáng lập thế kỷ 14. Nhấn mạnh giới luật, học thuật nghiêm ngặt. Tông phái của Đạt-lai Lạt-ma."],
]
add_table_custom(doc, ["Tông phái", "Đặc điểm cốt lõi"], sect_rows, col_widths=[1.4, 3.6])

add_heading_custom(doc, "7. Ẩm thực Tây Tạng — Các món ăn chính", level=2)
food_rows = [
    ["Tsampa", "Bột lúa mạch rang — lương thực chính, trộn với trà bơ hoặc sữa"],
    ["Momo", "Bánh bao/bánh xếp nhân thịt yak hoặc rau củ, hấp hoặc chiên"],
    ["Thịt yak khô", "Thịt yak phơi khô — đặc sản du mục, giàu dinh dưỡng"],
    ["Sữa chua Tây Tạng", "Làm từ sữa yak tươi, màu trắng sữa, vị chua đậm đà — đặc sản lễ hội Shoton"],
    ["Dresil", "Cơm ngọt — món tráng miệng trong các dịp đặc biệt"],
    ["Trà bơ (Po cha)", "Trà đun với nước soda, thêm muối và bơ sữa yak — nhiên liệu sinh tồn hàng ngày"],
    ["Rượu lúa mạch (Chang)", "Lên men từ lúa mạch cao nguyên, uống trong lễ hội, cưới xin"],
]
add_table_custom(doc, ["Món ăn", "Mô tả"], food_rows, col_widths=[1.4, 3.6])

add_heading_custom(doc, "8. Lịch sử cột mốc (tóm tắt nhanh cho HDV)", level=2)
hist_rows = [
    ["Thế kỷ 7", "Songtsen Gampo thống nhất Tây Tạng, xây Potala, Jokhang"],
    ["Thế kỷ 8", "Guru Rinpoche du nhập Mật tông, xây Samye"],
    ["792", "Đại biện luận Samye — Phật giáo Ấn Độ thắng Thiền tông Trung Hoa"],
    ["Thế kỷ 11–14", "Phát triển các tông phái: Nyingma, Kagyu, Sakya, Gelug"],
    ["Thế kỷ 14", "Tsongkhapa sáng lập Gelug (Hoàng giáo)"],
    ["Thế kỷ 17", "Đạt-lai Lạt-ma thứ 5 lên ngôi, xây Potala hiện tại"],
    ["1642–1950", "Thời kỳ chế độ thần quyền do Đạt-lai Lạt-ma cai trị"],
    ["1950", "Trận Chamdo — PLA tiến vào Tây Tạng"],
    ["1951", "Hiệp định 17 điểm — Tây Tạng chính thức thuộc Trung Quốc"],
    ["1959", "Nổi dậy Lhasa — Đạt-lai Lạt-ma thứ 14 lưu vong sang Ấn Độ"],
    ["1965", "Thành lập Khu tự trị Tây Tạng (TAR)"],
    ["1966–1976", "Cách mạng Văn hóa — hầu hết 6.000+ tu viện bị phá hủy"],
]
add_table_custom(doc, ["Thời kỳ", "Sự kiện"], hist_rows, col_widths=[1.0, 3.0])

add_heading_custom(doc, "9. Độ cao các thành phố chính", level=2)
city_rows = [
    ["Lhasa", "3.656m", "Thủ phủ cao nhất thế giới, 'Thành phố Ánh sáng'"],
    ["Shigatse", "3.840m", "Thành phố lớn thứ hai, tu viện Tashilhunpo"],
    ["Gyantse", "3.977m", "Cố đô lịch sử, tháp Kumbum"],
    ["Nagqu", "4.507m", "Thảo nguyên rộng lớn, lễ hội đua ngựa"],
    ["Ngari", "4.500m", "Vùng đất thiêng, núi Kailash"],
    ["Nyingchi", "3.000m", "Thấp nhất Tây Tạng, khí hậu ôn hòa nhất"],
    ["Chamdo", "3.256m", "Cửa ngõ phía đông"],
]
add_table_custom(doc, ["Thành phố", "Độ cao", "Đặc điểm"], city_rows, col_widths=[1.0, 0.8, 2.2])

doc.add_page_break()
