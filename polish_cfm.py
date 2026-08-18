# -*- coding: utf-8 -*-
from docx import Document


REPLACEMENTS = [
    ("TÂY TẠNG HUYỀN BÍ", "TÀI LIỆU VẬN HÀNH TOUR TÂY TẠNG"),
    ("CẨM NANG VẬN HÀNH & THUYẾT MINH HDV", "CẨM NANG VẬN HÀNH VÀ THUYẾT MINH"),
    ("Tour Tây Tạng Huyền Bí 9N8Đ", "Tour Tây Tạng 9N8Đ"),
    ("huyền bí", "tour"),
    ("Không phải bài giảng", "Trao đổi ngắn gọn"),
    ("không phải một kỳ nghỉ — nó là một cuộc hành trình biến đổi", "là một hành trình trải nghiệm"),
    ("Không liệt kê, chỉ có flow cảm xúc và layer tri thức.", "Ưu tiên mạch nội dung rõ ràng và dễ theo dõi."),
    ("flow cảm xúc", "mạch trải nghiệm"),
    ("Layer tri thức", "cấp độ nội dung"),
    ("Storytelling", "Thuyết minh"),
    ("MOT", "điểm nhấn quan trọng"),
    ("Moment of Truth", "điểm nhấn quan trọng"),
    ("Đừng xin lỗi. Kể chuyện.", "Giải thích ngắn gọn, đúng trọng tâm."),
    ("đừng xin lỗi", "giải thích ngắn gọn"),
    ("điểm check-in mang tính biểu tượng", "điểm tham quan tiêu biểu"),
    ("Đây không phải tour", "Đây là chương trình"),
    ("đỉnh cao", "điểm nhấn"),
    ("hùng vĩ", "ấn tượng"),
    ("huyền bí", "đặc trưng"),
    ("thánh địa", "điểm đến quan trọng"),
    ("thiêng liêng", "quan trọng về mặt văn hóa"),
    ("mang theo Tây Tạng", "tổng kết chuyến đi"),
    ("khoảnh khắc", "thời điểm"),
    ("chạm vào", "tiếp cận"),
    ("không thể tưởng tượng", "nội dung trọng tâm"),
    ("không thể", "khó"),
    ("siêu anh hùng", "biểu tượng"),
    ("nóc nhà", "vùng cao"),
    ("đặc quyền", "trải nghiệm"),
    ("điều chỉnh", "điều chỉnh phù hợp"),
    ("cát tường như ý", "chúc may mắn"),
    ("phiên bản đặc biệt", "bản rút gọn"),
    ("đậm chất", "mang tính"),
    ("rất đáng giá", "phù hợp"),
    ("rất khác biệt", "khác biệt"),
    ("vô cùng", "rất"),
    ("ấn tượng nhất", "quan trọng nhất"),
    ("hoàn toàn", "đủ"),
    ("hệ thống", "quy trình"),
    ("tuyệt đối", "không"),
    ("làm việc", "thực hiện"),
    ("cảm xúc", "trải nghiệm"),
    ("thần thoại", "truyền thuyết"),
    ("vị thần", "nhân vật"),
    ("nữ thần", "đỉnh Everest"),
    ("nếu có", "khi cần"),
    ("đánh dấu", "ghi nhận"),
    ("giữa muôn vàn sắc lam của đất trời", "trong khung cảnh cao nguyên"),
    ("không phải để đẹp", "để sử dụng"),
    ("nghe như", "được mô tả như"),
    ("chính là", "là"),
    ("ngay trước", "trước"),
    ("khác với", "không giống"),
    ("đủ tư cách", "phù hợp"),
    ("trân trọng", "ghi nhận"),
]


def replace_text(text: str) -> str:
    for old, new in REPLACEMENTS:
        text = text.replace(old, new)
    return text


def walk_paragraphs(doc):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = replace_text(run.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = replace_text(run.text)


doc = Document("CFM.docx")
walk_paragraphs(doc)
doc.save("CFM_polished.docx")

